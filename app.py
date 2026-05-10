import streamlit as st
import requests
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image,
                                 Table, TableStyle, HRFlowable, PageBreak)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT, TA_JUSTIFY
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, LineChart, Reference
import io

# =====================
# CONSTANTS DE COLOR
# =====================
MESOS_CA = {
    'January': 'Gener', 'February': 'Febrer', 'March': 'Marc',
    'April': 'Abril', 'May': 'Maig', 'June': 'Juny',
    'July': 'Juliol', 'August': 'Agost', 'September': 'Setembre',
    'October': 'Octubre', 'November': 'Novembre', 'December': 'Desembre'
}

# Colors PDF (reportlab)
C_BLAU = colors.HexColor('#1A3A5C')
C_BLAU_MIG = colors.HexColor('#2E86AB')
C_GROC = colors.HexColor('#F4A923')
C_GRIS_CLAR = colors.HexColor('#EEF2F7')
C_GRIS = colors.HexColor('#CCCCCC')
C_NEGRE = colors.HexColor('#1A1A1A')
C_VERD = colors.HexColor('#2E7D32')
C_BLANC = colors.white

# Colors hex (matplotlib, openpyxl)
H_BLAU = '#1A3A5C'
H_BLAU_MIG = '#2E86AB'
H_GROC = '#F4A923'
H_VERD = '#2E7D32'
H_GRIS = '#EEF2F7'

# Colors RGB (python-docx)
RGB_BLAU = RGBColor(0x1A, 0x3A, 0x5C)
RGB_BLAU_MIG = RGBColor(0x2E, 0x86, 0xAB)
RGB_GROC = RGBColor(0xF4, 0xA9, 0x23)
RGB_VERD = RGBColor(0x2E, 0x7D, 0x32)
RGB_BLANC = RGBColor(0xFF, 0xFF, 0xFF)
RGB_GRIS = RGBColor(0xEE, 0xF2, 0xF7)

# =====================
# CONFIGURACIÓ STREAMLIT
# =====================
st.set_page_config(page_title="EnergyDataGP", page_icon="⚡", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
    .header-bar {
        background: linear-gradient(135deg, #1A3A5C 0%, #2E86AB 100%);
        padding: 1.5rem 2rem; border-radius: 10px; margin-bottom: 1.5rem;
        box-shadow: 0 4px 15px rgba(26,58,92,0.3);
    }
    .header-title {
        color: white; font-size: 1.8rem; font-weight: 700;
        margin: 0; letter-spacing: 3px;
    }
    .header-sub { color: #F4A923; font-size: 0.9rem; margin: 0; margin-top: 6px; font-weight: 400; }
    .avis {
        background: #FFF8E1; border-left: 5px solid #F4A923;
        padding: 0.9rem 1.2rem; border-radius: 6px;
        font-size: 0.82rem; color: #555; margin-bottom: 1.2rem;
    }
    .seccio-titol {
        color: #1A3A5C; font-weight: 700; font-size: 1.05rem;
        border-bottom: 3px solid #F4A923; padding-bottom: 6px;
        margin-bottom: 1rem; margin-top: 1.5rem;
    }
    div[data-testid="metric-container"] {
        background: linear-gradient(135deg, #EEF2F7, #ffffff);
        border: 1px solid #CCCCCC; border-radius: 8px; padding: 0.5rem;
        border-top: 3px solid #1A3A5C;
    }
</style>
<div class="header-bar">
    <p class="header-title">⚡ ENERGYDATAGP</p>
    <p class="header-sub">Sistema de Generacio d'Informes Tecnics de Produccio Fotovoltaica</p>
</div>
<div class="avis">
    ⚠️ <b>Avis important:</b> Els informes generats tenen caracter orientatiu i no substitueixen
    la signatura d'un tecnic collegiat competent. Les dades provenen de PVGIS (Comissio Europea).
    EnergyDataGP no es fa responsable de les decisions preses sense validacio professional previa.
</div>
""", unsafe_allow_html=True)

# =====================
# FORMULARI
# =====================
st.markdown('<div class="seccio-titol">1. DADES DEL PROJECTE</div>', unsafe_allow_html=True)
col1, col2 = st.columns(2)
with col1:
    nom_projecte = st.text_input("Denominacio del projecte", "")
    num_expedient = st.text_input("Numero de referencia intern", "")
    nom_promotor = st.text_input("Promotor / Titular", "")
with col2:
    nom_tecnic = st.text_input("Tecnic responsable", "")
    num_collegio = st.text_input("Num. col·legiacio", "")
    adreca = st.text_input("Adreca de la instal·lacio", "")

st.markdown('<div class="seccio-titol">2. UBICACIO</div>', unsafe_allow_html=True)
col1, col2, col3 = st.columns(3)
with col1:
    lat = st.number_input("Latitud (N)", value=41.54, format="%.4f")
    lon = st.number_input("Longitud (E)", value=2.45, format="%.4f")
with col2:
    any_ref = st.selectbox("Any de referencia meteorologic", [2020, 2019, 2018, 2017])
with col3:
    st.info("Dades meteorologiques obtingudes automaticament de PVGIS (Comissio Europea)")

st.markdown('<div class="seccio-titol">3. PARAMETRES DE LA INSTAL·LACIO</div>', unsafe_allow_html=True)
mode_calc = st.radio(
    "Com vols definir la instal·lacio?",
    ["Per potencia total (kWp)", "Per superficie disponible (m2)", "Per nombre de panells"],
    horizontal=True
)
col1, col2 = st.columns(2)
with col2:
    st.markdown("**Parametres del panell fotovoltaic:**")
    potencia_panel_wp = st.number_input("Potencia pic del panell (Wp)", value=550, step=5)
    amplada_panel = st.number_input("Amplada del panell (m)", value=1.096, format="%.3f")
    llargada_panel = st.number_input("Llargada del panell (m)", value=2.384, format="%.3f")
    area_panel = amplada_panel * llargada_panel
    st.caption(f"Area per panell: {area_panel:.3f} m2")
with col1:
    if mode_calc == "Per potencia total (kWp)":
        potencia_kwp = st.number_input("Potencia pic total (kWp)", value=100.0)
        num_panels = int((potencia_kwp * 1000) / potencia_panel_wp)
        area_total = num_panels * area_panel
        area_coberta = area_total / 0.75
        st.metric("Nombre de panells", f"{num_panels}")
        st.metric("Area ocupada", f"{area_total:.1f} m2")
        st.metric("Area coberta necessaria (75%)", f"{area_coberta:.1f} m2")
    elif mode_calc == "Per superficie disponible (m2)":
        area_coberta = st.number_input("Superficie de coberta disponible (m2)", value=200.0)
        area_util = area_coberta * 0.75
        num_panels = int(area_util / area_panel)
        potencia_kwp = (num_panels * potencia_panel_wp) / 1000
        st.metric("Nombre de panells resultants", f"{num_panels}")
        st.metric("Potencia instal·lada resultant", f"{potencia_kwp:.1f} kWp")
        st.metric("Area util (75% coberta)", f"{area_util:.1f} m2")
    else:
        num_panels = st.number_input("Nombre de panells", value=136, step=1)
        potencia_kwp = (num_panels * potencia_panel_wp) / 1000
        area_total = num_panels * area_panel
        area_coberta = area_total / 0.75
        st.metric("Potencia instal·lada resultant", f"{potencia_kwp:.1f} kWp")
        st.metric("Area ocupada", f"{area_total:.1f} m2")
        st.metric("Area coberta necessaria (75%)", f"{area_coberta:.1f} m2")
perdues = st.number_input("Perdues del sistema (%)", value=14)

st.markdown('<div class="seccio-titol">4. PARAMETRES ECONOMICS</div>', unsafe_allow_html=True)
col1, col2, col3 = st.columns(3)
with col1:
    preu_kwh = st.number_input("Preu energia (EUR/kWh)", value=0.18, format="%.3f")
    consum_anual = st.number_input("Consum anual edifici (kWh/any)", value=96378)
with col2:
    cost_instalacio = st.number_input("Cost instal·lacio sense IVA (EUR, 0 = estima la IA)", value=0)
    iva = st.number_input("IVA (%)", value=21)
    if cost_instalacio > 0:
        cost_amb_iva = cost_instalacio * (1 + iva/100)
        st.metric("Cost amb IVA", f"{cost_amb_iva:,.0f} EUR")
    else:
        cost_amb_iva = 0
        st.info("La IA estimara el cost si es deixa en 0")
with col3:
    factor_co2 = st.number_input("Factor CO2 (kg CO2/kWh)", value=0.25, format="%.3f")
    preu_excedent = st.number_input("Preu excedent (EUR/kWh)", value=0.05, format="%.3f")

st.markdown('<div class="seccio-titol">5. DESGLOSSAMENT DEL PRESSUPOST</div>', unsafe_allow_html=True)
st.caption("Deixa el cost en blanc per als conceptes que vols que la IA estimi automaticament")
conceptes_default = [
    [f"Moduls FV {potencia_panel_wp} Wp", str(num_panels), ""],
    ["Inversors", "", ""], ["Estructura", "1", ""],
    ["Adequacio de sales tecniques", "", ""],
    ["Distribucio i proteccio electrica", "", ""],
    ["Mesura i punt de connexio", "", ""],
    ["Equips de control i monitoratge", "", ""],
    ["Mesures PRL", "", ""], ["Senyalitzacio i condicionament", "", ""],
    ["Gestio de residus", "", ""], ["Legalitzacio i projecte As-Built", "", ""],
    ["Seguretat i salut", "", ""], ["Transport i acopi de material", "", ""],
]
df_pressupost = pd.DataFrame(conceptes_default, columns=["Concepte", "Unitats", "Cost (EUR)"])
pressupost_editat = st.data_editor(
    df_pressupost, use_container_width=True, num_rows="dynamic",
    column_config={
        "Concepte": st.column_config.TextColumn("Concepte", width="large"),
        "Unitats": st.column_config.TextColumn("Unitats", width="small"),
        "Cost (EUR)": st.column_config.TextColumn("Cost (EUR)", width="medium"),
    }
)

st.markdown('<div class="seccio-titol">6. DESCRIPCIO DE L\'EMPLACAMENT</div>', unsafe_allow_html=True)
col1, col2 = st.columns(2)
with col1:
    tipus_coberta = st.selectbox("Tipus de coberta", [
        "Coberta plana", "Coberta inclinada", "Coberta coplanar orientada al Sud",
        "Coberta zenital", "Coberta mixta", "Facana", "Altres"])
    orientacio = st.selectbox("Orientacio principal", ["Sud", "Sud-Est", "Sud-Oest", "Est", "Oest", "Zenital"])
    inclinacio = st.number_input("Inclinacio (graus)", value=30)
with col2:
    tipus_activitat = st.text_input("Tipus d'activitat de l'edifici", "Administratiu")
    horari_activitat = st.selectbox("Horari principal d'activitat", [
        "Diurn (8h-18h)", "Nocturn (18h-6h)", "Continu (24h)", "Cap de setmana"])
    obstacles_ombres = st.selectbox("Presencia d'obstacles o ombres", [
        "No hi ha obstacles significatius",
        "Hi ha alguns obstacles menors",
        "Hi ha obstacles significatius"])

st.markdown("---")
with st.expander("📋 Avis legal i condicions d'us"):
    st.markdown("""
    **AVIS LEGAL — EnergyDataGP**

    Els documents generats no tenen valor juridic ni tecnic oficial i no substitueixen la memoria
    tecnica signada per un enginyer collegiat competent. Les dades provenen de PVGIS (Comissio Europea).
    EnergyDataGP no assumeix cap responsabilitat pels errors dels informes generats.
    L'us d'aquesta eina implica l'acceptacio d'aquestes condicions.
    """)

if st.button("⚡ GENERAR INFORME TECNIC", type="primary", use_container_width=True):
    if not nom_projecte or not nom_tecnic:
        st.error("Cal omplir com a minim la denominacio del projecte i el tecnic responsable.")
        st.stop()

    # =====================
    # PVGIS
    # =====================
    with st.spinner("Connectant amb PVGIS — Comissio Europea..."):
        url = "https://re.jrc.ec.europa.eu/api/v5_2/seriescalc"
        params = {
            "lat": lat, "lon": lon,
            "startyear": any_ref, "endyear": any_ref,
            "pvcalculation": 1, "peakpower": potencia_kwp,
            "loss": perdues, "outputformat": "json", "browser": 0
        }
        response = requests.get(url, params=params)
        df = pd.DataFrame(response.json()['outputs']['hourly'])
        df['time'] = pd.to_datetime(df['time'], format='%Y%m%d:%H%M')
        df.set_index('time', inplace=True)
        produccio_mensual = df['P'].resample('ME').sum() / 1000
        produccio_anual = produccio_mensual.sum()

    mes_max = MESOS_CA[produccio_mensual.idxmax().strftime('%B')]
    mes_min = MESOS_CA[produccio_mensual.idxmin().strftime('%B')]
    autoconsum_ratio = min(consum_anual / produccio_anual, 1.0) if produccio_anual > 0 else 0.5
    autoconsum_kwh = produccio_anual * autoconsum_ratio
    excedents_kwh = produccio_anual - autoconsum_kwh
    estalvi_autoconsum = autoconsum_kwh * preu_kwh
    benefici_excedent = excedents_kwh * preu_excedent
    estalvi_anual = estalvi_autoconsum + benefici_excedent
    co2_estalviat = produccio_anual * factor_co2

    # =====================
    # IA — GROQ
    # =====================
    with st.spinner("Generant textos i pressupost amb IA..."):
        try:
            from groq import Groq
            client = Groq(api_key=st.secrets["GROQ_API_KEY"])

            # PRESSUPOST
            costos_buits = pressupost_editat['Cost (EUR)'].apply(
                lambda x: str(x).strip() == '' or str(x).strip() == 'None').any()
            if costos_buits:
                conceptes_str = ""
                for _, row in pressupost_editat.iterrows():
                    concepte = str(row['Concepte']) if row['Concepte'] else ''
                    unitats = str(row['Unitats']) if row['Unitats'] else ''
                    cost = str(row['Cost (EUR)']) if row['Cost (EUR)'] else ''
                    if concepte:
                        estat = f"COST CONEGUT: {cost} EUR" if cost and cost not in ['None', ''] else "ESTIMA EL COST"
                        conceptes_str += f"- {concepte} (Unitats: {unitats if unitats else 'N/A'}): {estat}\n"

                prompt_pressupost = f"""Ets un enginyer expert en instal·lacions fotovoltaiques a Catalunya.
Estima els costos dels conceptes marcats com "ESTIMA EL COST" per a una instal·lacio de {potencia_kwp:.1f} kWp amb {num_panels} panells de {potencia_panel_wp} Wp.

Preus de referencia del sector a Catalunya:
- Moduls FV 550 Wp: 120 EUR/unitat
- Inversors: 150-200 EUR/kW
- Estructura suport: 70-90 EUR/modul
- Adequacio sales tecniques: 3.000-5.000 EUR
- Distribucio i proteccio electrica: 3.500-5.500 EUR
- Mesura i punt de connexio: 1.500-2.000 EUR
- Equips control i monitoratge: 1.000-1.500 EUR
- Mesures PRL: 4.000-6.000 EUR
- Senyalitzacio i condicionament: 1.000-1.500 EUR
- Gestio de residus: 1.800-2.500 EUR
- Legalitzacio i projecte As-Built: 1.800-2.200 EUR
- Seguretat i salut: 1.500-2.200 EUR
- Transport i acopi: 1.800-2.500 EUR

Conceptes del pressupost:
{conceptes_str}

Respon NOMES en format JSON sense cap text addicional:
{{"pressupost": [{{"concepte": "nom", "unitats": "valor", "cost": 1234.00}}, ...]}}

Inclou TOTS els conceptes. Per als que ja tenien cost conegut, usa el cost indicat."""

                resp_pres = client.chat.completions.create(
                    messages=[{"role": "user", "content": prompt_pressupost}],
                    model="llama-3.3-70b-versatile",
                )
                import json
                text_resp = resp_pres.choices[0].message.content.strip()
                if '```' in text_resp:
                    text_resp = text_resp.split('```')[1]
                    if text_resp.startswith('json'):
                        text_resp = text_resp[4:]
                pressupost_ia = json.loads(text_resp)
                pressupost_final = pressupost_ia['pressupost']
            else:
                pressupost_final = []
                for _, row in pressupost_editat.iterrows():
                    if row['Concepte']:
                        try:
                            cost_val = float(str(row['Cost (EUR)']).replace('.', '').replace(',', '.').replace('EUR', '').strip())
                        except:
                            cost_val = 0
                        pressupost_final.append({'concepte': str(row['Concepte']), 'unitats': str(row['Unitats']) if row['Unitats'] else '', 'cost': cost_val})

            total_sense_iva = sum(item['cost'] for item in pressupost_final)
            total_amb_iva = total_sense_iva * (1 + iva/100)
            if cost_instalacio == 0:
                cost_amb_iva = total_amb_iva
            anys_retorn = cost_amb_iva / estalvi_anual if estalvi_anual > 0 else 0

            # TEXT ESTAT ACTUAL
            resp_estat = client.chat.completions.create(
                messages=[{"role": "user", "content": f"""Ets un enginyer expert en energia solar redactant un informe tecnic en catala formal.
Redacta 3-4 frases d'ESTAT ACTUAL de l'edifici sense titol:
- Coberta: {tipus_coberta}, orientacio {orientacio}, inclinacio {inclinacio} graus
- Activitat: {tipus_activitat}, horari {horari_activitat}
- Obstacles: {obstacles_ombres}
Comenca directament amb la descripcio. Usa un registre tecnic formal."""}],
                model="llama-3.3-70b-versatile",
            )
            text_estat = resp_estat.choices[0].message.content

            # TEXT PROPOSTA
            resp_proposta = client.chat.completions.create(
                messages=[{"role": "user", "content": f"""Ets un enginyer expert en energia solar redactant un informe tecnic en catala formal.
Redacta 3-4 frases de PROPOSTA DE MILLORA sense titol:
- {num_panels} moduls de {potencia_panel_wp} Wp, potencia total {potencia_kwp:.1f} kWp
- Orientacio {orientacio}, inclinacio {inclinacio} graus
- Produccio anual estimada: {produccio_anual:,.0f} kWh
- {obstacles_ombres}
Comenca directament amb la proposta. Usa un registre tecnic formal."""}],
                model="llama-3.3-70b-versatile",
            )
            text_proposta = resp_proposta.choices[0].message.content

            # TEXT CONCLUSIONS
            resp_conclusions = client.chat.completions.create(
                messages=[{"role": "user", "content": f"""Ets un enginyer expert en energia solar redactant un informe tecnic en catala formal.
Redacta 4-5 frases de CONCLUSIONS sense titol:
- Potencia: {potencia_kwp:.1f} kWp, {num_panels} panells de {potencia_panel_wp} Wp
- Produccio anual: {produccio_anual:,.0f} kWh
- Autoconsum: {autoconsum_kwh:,.0f} kWh/any, Excedents: {excedents_kwh:,.0f} kWh/any
- Cost amb IVA: {cost_amb_iva:,.0f} EUR
- Estalvi anual: {estalvi_anual:,.0f} EUR/any, Payback: {anys_retorn:.1f} anys
- CO2 estalviat: {co2_estalviat:,.0f} kg/any
Comenca directament amb les conclusions. Usa un registre tecnic formal."""}],
                model="llama-3.3-70b-versatile",
            )
            text_conclusions = resp_conclusions.choices[0].message.content
            st.success("✅ Textos i pressupost generats per IA correctament")

        except Exception as e:
            st.warning(f"⚠️ IA no disponible ({e}). S'usa text estandard.")
            pressupost_final = []
            for _, row in pressupost_editat.iterrows():
                if row['Concepte']:
                    try:
                        cost_val = float(str(row['Cost (EUR)']).replace('.', '').replace(',', '.').replace('EUR', '').strip())
                    except:
                        cost_val = 0
                    pressupost_final.append({'concepte': str(row['Concepte']), 'unitats': str(row['Unitats']) if row['Unitats'] else '', 'cost': cost_val})
            total_sense_iva = sum(item['cost'] for item in pressupost_final)
            total_amb_iva = total_sense_iva * (1 + iva/100)
            if cost_instalacio == 0:
                cost_amb_iva = total_amb_iva
            anys_retorn = cost_amb_iva / estalvi_anual if estalvi_anual > 0 else 0
            text_estat = f"L'edifici disposa d'una {tipus_coberta.lower()} amb orientacio {orientacio} i inclinacio de {inclinacio} graus. L'activitat principal es {tipus_activitat.lower()} amb horari {horari_activitat.lower()}. {obstacles_ombres}."
            text_proposta = f"Es proposa instal·lar {num_panels} moduls fotovoltaics de {potencia_panel_wp} Wp, amb una potencia total de {potencia_kwp:.1f} kWp i orientacio {orientacio} a {inclinacio} graus d'inclinacio. La produccio anual estimada es de {produccio_anual:,.0f} kWh."
            text_conclusions = f"La instal·lacio fotovoltaica {nom_projecte} presenta una produccio anual estimada de {produccio_anual:,.0f} kWh, un estalvi economic de {estalvi_anual:,.0f} EUR/any i un retorn de la inversio en {anys_retorn:.1f} anys. La instal·lacio contribuira a la reduccio d'emissions de CO2 en {co2_estalviat:,.0f} kg per any."

    # =====================
    # GRAFIQUES MATPLOTLIB
    # =====================
    with st.spinner("Generant grafiques tecniques..."):

        plt.rcParams['font.family'] = 'DejaVu Sans'
        plt.rcParams['axes.spines.top'] = False
        plt.rcParams['axes.spines.right'] = False

        # G1 — Produccio mensual
        fig1, ax1 = plt.subplots(figsize=(13, 4.5))
        bars = ax1.bar(range(12), produccio_mensual.values, color=H_BLAU, alpha=0.85, width=0.65, zorder=3)
        ax1.set_xticks(range(12))
        ax1.set_xticklabels(['Gen','Feb','Mar','Abr','Mai','Jun','Jul','Ago','Set','Oct','Nov','Des'], fontsize=10)
        ax1.set_ylabel('Energia produida (kWh)', fontsize=10, color=H_BLAU)
        ax1.set_title(f'Figura 1. Produccio energetica mensual — {nom_projecte}', fontsize=11, fontweight='bold', color=H_BLAU, pad=12)
        mitj = produccio_mensual.mean()
        ax1.axhline(mitj, color=H_GROC, linestyle='--', linewidth=2, label=f'Mitjana: {mitj:,.0f} kWh', zorder=4)
        for i, val in enumerate(produccio_mensual.values):
            ax1.text(i, val + produccio_mensual.max()*0.02, f'{val:,.0f}', ha='center', fontsize=7.5, color=H_BLAU, fontweight='500')
        ax1.legend(fontsize=9, loc='upper left')
        ax1.grid(axis='y', alpha=0.3, zorder=0)
        ax1.set_facecolor('#FAFBFC')
        fig1.patch.set_facecolor('white')
        # Linia inferior de color
        ax1.axhline(0, color=H_GROC, linewidth=3)
        plt.tight_layout()
        plt.savefig('g1_mensual.png', dpi=180, bbox_inches='tight', facecolor='white')
        plt.close()

        # G2 — Corba estiu vs hivern
        dia_estiu = df.loc[f'{any_ref}-06-21', 'P'] / 1000
        dia_hivern = df.loc[f'{any_ref}-12-21', 'P'] / 1000
        fig2, ax2 = plt.subplots(figsize=(13, 4.5))
        ax2.plot(range(len(dia_estiu)), dia_estiu.values, color=H_BLAU_MIG, linewidth=2.5, label="Solstici d'estiu (21 juny)", zorder=3)
        ax2.plot(range(len(dia_hivern)), dia_hivern.values, color=H_BLAU, linewidth=2.5, label="Solstici d'hivern (21 desembre)", zorder=3)
        ax2.fill_between(range(len(dia_estiu)), dia_estiu.values, alpha=0.12, color=H_BLAU_MIG)
        ax2.fill_between(range(len(dia_hivern)), dia_hivern.values, alpha=0.12, color=H_BLAU)
        ax2.set_ylabel('Potencia generada (kW)', fontsize=10, color=H_BLAU)
        ax2.set_title('Figura 2. Corba de generacio horaria — Comparativa estacional', fontsize=11, fontweight='bold', color=H_BLAU, pad=12)
        ax2.legend(fontsize=9)
        ax2.grid(alpha=0.3, zorder=0)
        ax2.set_facecolor('#FAFBFC')
        ax2.axhline(0, color=H_GROC, linewidth=3)
        plt.tight_layout()
        plt.savefig('g2_corbes.png', dpi=180, bbox_inches='tight', facecolor='white')
        plt.close()

        # G3 — Corbes PVsyst per mesos tipus
        mesos_tipics = {'Gener': f'{any_ref}-01-15', 'Abril': f'{any_ref}-04-15', 'Juliol': f'{any_ref}-07-15', 'Octubre': f'{any_ref}-10-15'}
        colors_m = [H_BLAU, H_BLAU_MIG, H_GROC, H_VERD]
        fig3, ax3 = plt.subplots(figsize=(13, 5))
        for (mes_nom, data), color in zip(mesos_tipics.items(), colors_m):
            try:
                dia = df.loc[data, 'P'] / 1000
                ax3.plot(range(len(dia)), dia.values, linewidth=2.2, label=mes_nom, color=color, zorder=3)
                ax3.fill_between(range(len(dia)), dia.values, alpha=0.06, color=color)
            except:
                pass
        ax3.set_xlabel('Hora del dia', fontsize=10)
        ax3.set_ylabel('Potencia generada (kW)', fontsize=10, color=H_BLAU)
        ax3.set_title('Figura 3. Corbes de generacio horaria per mesos tipus (estil PVsyst)', fontsize=11, fontweight='bold', color=H_BLAU, pad=12)
        ax3.set_xticks(range(0, 24, 2))
        ax3.set_xticklabels([f'{h:02d}:00' for h in range(0, 24, 2)], rotation=45, fontsize=8)
        ax3.legend(fontsize=9, ncol=4, loc='upper center')
        ax3.grid(alpha=0.3, zorder=0)
        ax3.set_facecolor('#FAFBFC')
        ax3.axhline(0, color=H_GROC, linewidth=3)
        plt.tight_layout()
        plt.savefig('g3_horaria.png', dpi=180, bbox_inches='tight', facecolor='white')
        plt.close()

        # G4 — ROI
        anys_list = list(range(0, 26))
        flux_acumulat = [-cost_amb_iva + estalvi_anual * a for a in anys_list]
        fig4, ax4 = plt.subplots(figsize=(13, 4.5))
        colors_bars = [H_BLAU_MIG if v < 0 else H_VERD for v in flux_acumulat]
        ax4.bar(anys_list, flux_acumulat, color=colors_bars, alpha=0.85, width=0.7, zorder=3)
        ax4.axhline(0, color=H_BLAU, linewidth=1.5, zorder=4)
        ax4.set_xlabel('Anys des de la posada en marxa', fontsize=10)
        ax4.set_ylabel('Flux de caixa acumulat (EUR)', fontsize=10, color=H_BLAU)
        ax4.set_title('Figura 4. Analisi del retorn de la inversio (ROI)', fontsize=11, fontweight='bold', color=H_BLAU, pad=12)
        ax4.grid(axis='y', alpha=0.3, zorder=0)
        ax4.set_facecolor('#FAFBFC')
        llegenda = [mpatches.Patch(color=H_BLAU_MIG, label='Periode de retorn'), mpatches.Patch(color=H_VERD, label='Benefici net')]
        ax4.legend(handles=llegenda, fontsize=9)
        ax4.axhline(0, color=H_GROC, linewidth=3)
        plt.tight_layout()
        plt.savefig('g4_roi.png', dpi=180, bbox_inches='tight', facecolor='white')
        plt.close()

        # G5 — Autoconsum pie
        fig5, ax5 = plt.subplots(figsize=(7, 5))
        sizes = [autoconsum_kwh, excedents_kwh]
        labels_pie = [f'Autoconsum\n{autoconsum_kwh:,.0f} kWh', f'Excedents exportats\n{excedents_kwh:,.0f} kWh']
        colors_pie = [H_BLAU, H_BLAU_MIG]
        wedges, texts, autotexts = ax5.pie(sizes, labels=labels_pie, colors=colors_pie,
               autopct='%1.1f%%', startangle=90, textprops={'fontsize': 10},
               wedgeprops={'edgecolor': 'white', 'linewidth': 2})
        for at in autotexts:
            at.set_color('white')
            at.set_fontweight('bold')
        ax5.set_title("Figura 5. Distribucio de l'energia generada", fontsize=11, fontweight='bold', color=H_BLAU, pad=12)
        plt.tight_layout()
        plt.savefig('g5_autoconsum.png', dpi=180, bbox_inches='tight', facecolor='white')
        plt.close()

    # Resum web
    st.success("✅ Informe generat correctament")
    st.markdown("### Resum executiu")
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Produccio anual", f"{produccio_anual:,.0f} kWh")
    c2.metric("Autoconsum", f"{autoconsum_kwh:,.0f} kWh")
    c3.metric("Estalvi anual", f"{estalvi_anual:,.0f} EUR")
    c4.metric("Payback", f"{anys_retorn:.1f} anys")
    c5.metric("CO2 estalviat", f"{co2_estalviat:,.0f} kg")
    st.pyplot(fig1)
    st.pyplot(fig3)

    # =====================
    # HELPERS PDF
    # =====================
    def estil_taula_standard():
        return TableStyle([
            ('BACKGROUND', (0,0), (-1,0), C_BLAU),
            ('TEXTCOLOR', (0,0), (-1,0), C_BLANC),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 9),
            ('LINEBELOW', (0,0), (-1,0), 2, C_GROC),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [C_GRIS_CLAR, C_BLANC]),
            ('FONTSIZE', (0,1), (-1,-1), 9),
            ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
            ('GRID', (0,0), (-1,-1), 0.3, C_GRIS),
            ('PADDING', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ])

    def seccio_pdf(elements, num, titol, estil_s):
        elements.append(Spacer(1, 0.3*cm))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=C_BLAU_MIG, spaceAfter=0))
        elements.append(Paragraph(f"{num}. {titol}", estil_s))
        elements.append(HRFlowable(width="100%", thickness=3, color=C_GROC, spaceAfter=8))

    # =====================
    # PDF
    # =====================
    with st.spinner("Composant el document PDF..."):
        doc_pdf = SimpleDocTemplate(
            "informe_energydatagp.pdf", pagesize=A4,
            rightMargin=2.2*cm, leftMargin=2.2*cm, topMargin=2*cm, bottomMargin=2.5*cm
        )

        styles = getSampleStyleSheet()

        # Estils PDF
        s_seccio = ParagraphStyle('seccio', fontSize=12, fontName='Helvetica-Bold',
            textColor=C_BLAU, spaceBefore=8, spaceAfter=4)
        s_normal = ParagraphStyle('normal', fontSize=9, fontName='Helvetica',
            textColor=C_NEGRE, leading=14, spaceAfter=4, alignment=TA_JUSTIFY)
        s_peu = ParagraphStyle('peu', fontSize=7, fontName='Helvetica',
            textColor=colors.grey, alignment=TA_CENTER)
        s_avis = ParagraphStyle('avis', fontSize=7.5, fontName='Helvetica',
            textColor=colors.grey, leading=11, spaceAfter=6)
        s_cap_taula = ParagraphStyle('cap_taula', fontSize=8, fontName='Helvetica-Bold',
            textColor=C_BLANC)
        s_val_taula = ParagraphStyle('val_taula', fontSize=8.5, fontName='Helvetica',
            textColor=C_NEGRE)

        elements = []

        # ---- CAPÇALERA PDF ----
        elements.append(HRFlowable(width="100%", thickness=6, color=C_BLAU, spaceAfter=0))
        cap_data = [[
            Paragraph('<b><font size=15 color="#1A3A5C">ENERGYDATAGP</font></b><br/>'
                     '<font size=8 color="#2E86AB">Solucions d\'Analisi de Dades Energetiques</font>',
                     ParagraphStyle('cap', fontSize=9, fontName='Helvetica-Bold', textColor=C_BLAU)),
            Paragraph('<b><font size=12 color="#1A3A5C">INFORME TECNIC DE PRODUCCIO</font></b><br/>'
                     '<font size=10 color="#1A3A5C">INSTAL·LACIO FOTOVOLTAICA</font>',
                     ParagraphStyle('cap2', fontSize=12, fontName='Helvetica-Bold', textColor=C_BLAU, alignment=TA_CENTER)),
            Paragraph(f'<font size=8 color="#1A3A5C">Ref.: <b>{num_expedient if num_expedient else "-"}</b></font><br/>'
                     f'<font size=8 color="#1A3A5C">Data: <b>{datetime.now().strftime("%d/%m/%Y")}</b></font><br/>'
                     f'<font size=8 color="#1A3A5C">Versio: 1.0</font>',
                     ParagraphStyle('cap3', fontSize=9, fontName='Helvetica', textColor=C_BLAU, alignment=TA_RIGHT)),
        ]]
        taula_cap = Table(cap_data, colWidths=[5.5*cm, 8*cm, 4*cm])
        taula_cap.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F0F4F8')),
            ('PADDING', (0,0), (-1,-1), 12),
        ]))
        elements.append(taula_cap)
        elements.append(HRFlowable(width="100%", thickness=4, color=C_GROC, spaceAfter=16))

        # ---- SECCIÓ 1 — IDENTIFICACIÓ ----
        seccio_pdf(elements, 1, "IDENTIFICACIO DEL PROJECTE", s_seccio)
        dades_id = [
            [Paragraph('<b>Denominacio:</b>', s_cap_taula), Paragraph(nom_projecte, s_val_taula),
             Paragraph('<b>Promotor:</b>', s_cap_taula), Paragraph(nom_promotor if nom_promotor else '-', s_val_taula)],
            [Paragraph('<b>Referencia:</b>', s_cap_taula), Paragraph(num_expedient if num_expedient else '-', s_val_taula),
             Paragraph('<b>Adreca:</b>', s_cap_taula), Paragraph(adreca if adreca else '-', s_val_taula)],
            [Paragraph('<b>Coordenades:</b>', s_cap_taula), Paragraph(f'{lat}N, {lon}E', s_val_taula),
             Paragraph('<b>Any meteorologic:</b>', s_cap_taula), Paragraph(str(any_ref), s_val_taula)],
            [Paragraph('<b>Tecnic responsable:</b>', s_cap_taula), Paragraph(nom_tecnic, s_val_taula),
             Paragraph('<b>Col·legiacio:</b>', s_cap_taula), Paragraph(num_collegio if num_collegio else '-', s_val_taula)],
        ]
        taula_id = Table(dades_id, colWidths=[3.5*cm, 6.5*cm, 3.5*cm, 4*cm])
        ts_id = estil_taula_standard()
        ts_id.add('BACKGROUND', (0,0), (-1,-1), C_GRIS_CLAR)
        ts_id.add('BACKGROUND', (0,0), (0,-1), C_BLAU)
        ts_id.add('BACKGROUND', (2,0), (2,-1), C_BLAU)
        ts_id.add('ROWBACKGROUNDS', (0,0), (-1,-1), [C_GRIS_CLAR, C_BLANC])
        taula_id.setStyle(ts_id)
        elements.append(taula_id)

        # ---- SECCIÓ 2 — ESTAT ACTUAL ----
        seccio_pdf(elements, 2, "ESTAT ACTUAL", s_seccio)
        elements.append(Paragraph(text_estat, s_normal))

        # ---- SECCIÓ 3 — PROPOSTA ----
        seccio_pdf(elements, 3, "PROPOSTA DE MILLORA", s_seccio)
        elements.append(Paragraph(text_proposta, s_normal))

        # ---- SECCIÓ 4 — PRESSUPOST ----
        seccio_pdf(elements, 4, "PRESSUPOST DETALLAT", s_seccio)
        pres_data = [['CONCEPTE', 'UNITATS', 'COST']]
        for item in pressupost_final:
            cost_str = f"{item['cost']:,.2f} EUR" if item['cost'] > 0 else '-'
            pres_data.append([item['concepte'], str(item['unitats']), cost_str])
        pres_data.append(['TOTAL (sense IVA)', '', f'{total_sense_iva:,.2f} EUR'])
        pres_data.append([f'TOTAL (amb IVA {iva}%)', '', f'{total_amb_iva:,.2f} EUR'])
        taula_pres = Table(pres_data, colWidths=[10*cm, 2.5*cm, 5*cm])
        ts_pres = estil_taula_standard()
        ts_pres.add('BACKGROUND', (0,-2), (-1,-1), colors.HexColor('#D6E4F0'))
        ts_pres.add('FONTNAME', (0,-2), (-1,-1), 'Helvetica-Bold')
        ts_pres.add('TEXTCOLOR', (0,-2), (-1,-1), C_BLAU)
        ts_pres.add('LINEABOVE', (0,-2), (-1,-2), 1.5, C_BLAU)
        ts_pres.add('ALIGN', (2,0), (2,-1), 'RIGHT')
        taula_pres.setStyle(ts_pres)
        elements.append(taula_pres)

        # ---- SECCIÓ 5 — PARAMETRES TECNICS ----
        seccio_pdf(elements, 5, "PARAMETRES TECNICS DE LA INSTAL·LACIO", s_seccio)
        dades_tec = [
            ['Parametre', 'Valor', 'Unitat', 'Observacions'],
            ['Potencia pic instal·lada', f'{potencia_kwp:.1f}', 'kWp', 'Condicions estandard (STC)'],
            ['Nombre de moduls', f'{num_panels}', 'ut.', f'{potencia_panel_wp} Wp per modul'],
            ['Dimensions modul', f'{llargada_panel:.3f}x{amplada_panel:.3f}', 'm', f'Area: {area_panel:.3f} m2/modul'],
            ['Perdues del sistema', f'{perdues}', '%', 'Cablejat, inversor, bruticia, temperatura'],
            ['Produccio anual estimada', f'{produccio_anual:,.0f}', 'kWh/any', f'Any de referencia: {any_ref}'],
            ['Mes de maxima produccio', mes_max, f'{produccio_mensual.max():,.0f} kWh', 'Pic de produccio estival'],
            ['Mes de minima produccio', mes_min, f'{produccio_mensual.min():,.0f} kWh', 'Minim de produccio hivernal'],
            ['Hores equivalents de sol', f'{produccio_anual/potencia_kwp:,.0f}', 'HES/any', 'Hores a potencia nominal'],
            ['Orientacio / Inclinacio', orientacio, f'{inclinacio} graus', tipus_coberta],
        ]
        taula_tec = Table(dades_tec, colWidths=[5.5*cm, 3*cm, 3*cm, 6*cm])
        taula_tec.setStyle(estil_taula_standard())
        elements.append(taula_tec)

        # ---- SECCIÓ 6 — PRODUCCIO MENSUAL ----
        seccio_pdf(elements, 6, "ANALISI DE LA PRODUCCIO ENERGETICA MENSUAL", s_seccio)
        elements.append(Image('g1_mensual.png', width=16.5*cm, height=6.2*cm))
        elements.append(Spacer(1, 0.2*cm))
        elements.append(Paragraph(
            f"La instal·lacio fotovoltaica presenta una produccio anual estimada de "
            f"<b>{produccio_anual:,.0f} kWh</b>, d'acord amb les dades meteorologiques de l'any {any_ref} "
            f"de PVGIS. El mes de maxima produccio correspon a <b>{mes_max}</b> "
            f"({produccio_mensual.max():,.0f} kWh), mentre que el mes de minima produccio es "
            f"<b>{mes_min}</b> ({produccio_mensual.min():,.0f} kWh).", s_normal))

        # ---- SECCIÓ 7 — CORBES HORARIES ----
        seccio_pdf(elements, 7, "CORBES DE GENERACIO HORARIA", s_seccio)
        elements.append(Image('g2_corbes.png', width=16.5*cm, height=6.2*cm))
        elements.append(Spacer(1, 0.2*cm))
        elements.append(Image('g3_horaria.png', width=16.5*cm, height=6.2*cm))
        elements.append(Spacer(1, 0.2*cm))
        elements.append(Paragraph(
            f"La figura 2 mostra la corba de generacio horaria per als solsticis d'estiu i d'hivern. "
            f"La figura 3 representa les corbes per als mesos tipus representatius de l'any, seguint "
            f"la metodologia habitual en estudis de produccio fotovoltaica. "
            f"La potencia maxima durant el solstici d'estiu es de <b>{dia_estiu.max():.1f} kW</b>, "
            f"mentre que durant el solstici d'hivern es de <b>{dia_hivern.max():.1f} kW</b>.", s_normal))

        # ---- SECCIÓ 8 — ANALISI ECONOMICA ----
        seccio_pdf(elements, 8, "ANALISI ECONOMICA I RETORN DE LA INVERSIO", s_seccio)

        # Taula simulació
        dades_sim = [
            ['', 'Consum xarxa (kWh/any)', 'Generacio FV (kWh/any)', 'Autoconsum (kWh/any)', 'Excedents (kWh/any)'],
            ['Energia', f'{consum_anual:,.0f}', f'{produccio_anual:,.0f}', f'{autoconsum_kwh:,.0f}', f'{excedents_kwh:,.0f}'],
            ['Cost / Benefici', '—', '—', f'{estalvi_autoconsum:,.0f} EUR', f'{benefici_excedent:,.0f} EUR'],
        ]
        taula_sim = Table(dades_sim, colWidths=[2.5*cm, 3.5*cm, 3.5*cm, 3.5*cm, 3.5*cm])
        ts_sim = estil_taula_standard()
        ts_sim.add('FONTNAME', (0,1), (0,-1), 'Helvetica-Bold')
        ts_sim.add('TEXTCOLOR', (0,1), (0,-1), C_BLAU)
        taula_sim.setStyle(ts_sim)
        elements.append(taula_sim)
        elements.append(Spacer(1, 0.3*cm))

        # Taula resum inversió
        dades_eco = [
            ['Consum actual\n(kWh/any)', 'Estalvi energia\n(kWh/any)', 'Estalvi economic\n(EUR/any)', 'Inversio\n(EUR)', 'Payback\n(anys)', 'Estalvi CO2\n(kg/any)'],
            [f'{consum_anual:,.0f}', f'{autoconsum_kwh:,.0f}', f'{estalvi_anual:,.0f}', f'{cost_amb_iva:,.0f}', f'{anys_retorn:.1f}', f'{co2_estalviat:,.0f}'],
        ]
        taula_eco = Table(dades_eco, colWidths=[3*cm, 3*cm, 3*cm, 2.5*cm, 2.5*cm, 3.5*cm])
        ts_eco = estil_taula_standard()
        ts_eco.add('ALIGN', (0,0), (-1,-1), 'CENTER')
        taula_eco.setStyle(ts_eco)
        elements.append(taula_eco)
        elements.append(Spacer(1, 0.3*cm))

        # Grafiques economia
        col_g = [[Image('g5_autoconsum.png', width=8*cm, height=6*cm),
                   Image('g4_roi.png', width=9*cm, height=6*cm)]]
        taula_g = Table(col_g, colWidths=[8.5*cm, 9*cm])
        taula_g.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('PADDING', (0,0), (-1,-1), 4)]))
        elements.append(taula_g)

        # ---- SECCIÓ 9 — CONCLUSIONS ----
        seccio_pdf(elements, 9, "CONCLUSIONS", s_seccio)
        elements.append(Paragraph(text_conclusions, s_normal))
        elements.append(Spacer(1, 0.8*cm))

        # Signatura
        sig_data = [
            [Paragraph('<b>El/La tecnic/a responsable</b>', ParagraphStyle('sig', fontSize=9, fontName='Helvetica-Bold', textColor=C_BLAU)),
             Paragraph('<b>Conforme el/la promotor/a</b>', ParagraphStyle('sig2', fontSize=9, fontName='Helvetica-Bold', textColor=C_BLAU))],
            [Paragraph(f'<br/><br/><br/>Nom: {nom_tecnic}<br/>Col·legiacio: {num_collegio if num_collegio else "—"}<br/>Data: {datetime.now().strftime("%d/%m/%Y")}',
                      ParagraphStyle('sig3', fontSize=9, fontName='Helvetica', textColor=C_NEGRE, leading=14)),
             Paragraph('<br/><br/><br/>Nom:<br/>DNI/NIF:<br/>Data:',
                      ParagraphStyle('sig4', fontSize=9, fontName='Helvetica', textColor=C_NEGRE, leading=14))],
        ]
        taula_sig = Table(sig_data, colWidths=[8.5*cm, 8.5*cm])
        taula_sig.setStyle(TableStyle([
            ('BOX', (0,0), (0,-1), 0.8, C_BLAU), ('BOX', (1,0), (1,-1), 0.8, C_BLAU),
            ('LINEBELOW', (0,0), (0,0), 2, C_GROC), ('LINEBELOW', (1,0), (1,0), 2, C_GROC),
            ('BACKGROUND', (0,0), (0,0), C_GRIS_CLAR), ('BACKGROUND', (1,0), (1,0), C_GRIS_CLAR),
            ('PADDING', (0,0), (-1,-1), 8),
        ]))
        elements.append(taula_sig)
        elements.append(Spacer(1, 0.6*cm))

        # Peu PDF
        elements.append(HRFlowable(width="100%", thickness=4, color=C_BLAU, spaceAfter=4))
        peu_data = [[
            Paragraph('<b>EnergyDataGP</b> — Solucions d\'Analisi de Dades Energetiques',
                     ParagraphStyle('pef', fontSize=8, fontName='Helvetica', textColor=C_BLAU)),
            Paragraph(f'Informe generat el {datetime.now().strftime("%d/%m/%Y")} | PVGIS © Comissio Europea',
                     ParagraphStyle('pec', fontSize=7, fontName='Helvetica', textColor=colors.grey, alignment=TA_CENTER)),
            Paragraph('www.energydatagp.com',
                     ParagraphStyle('ped', fontSize=8, fontName='Helvetica', textColor=C_BLAU_MIG, alignment=TA_RIGHT)),
        ]]
        taula_peu = Table(peu_data, colWidths=[6*cm, 7*cm, 4.5*cm])
        taula_peu.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('PADDING', (0,0), (-1,-1), 4)]))
        elements.append(taula_peu)
        elements.append(Spacer(1, 0.2*cm))
        elements.append(Paragraph(
            "<b>AVIS LEGAL:</b> Document generat amb EnergyDataGP amb finalitat orientativa. "
            "Dades de produccio: PVGIS © Comissio Europea. No substitueix la memoria tecnica "
            "oficial signada per un tecnic col·legiat competent. EnergyDataGP no es fa responsable "
            "de les decisions preses a partir d'aquest document sense validacio professional previa.",
            s_avis))

        doc_pdf.build(elements)

    # =====================
    # WORD
    # =====================
    with st.spinner("Generant document Word..."):
        doc_word = Document()

        # Configuració de pàgina A4
        from docx.shared import Mm
        section = doc_word.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

        # Helpers Word
        def set_cell_bg(cell, hex_color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)

        def set_cell_border_bottom(cell, hex_color, size=24):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(size))
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), hex_color)
            tcBorders.append(bottom)
            tcPr.append(tcBorders)

        def add_heading_word(doc, text, level=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            # Linia superior blava
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '4')
            top.set(qn('w:space'), '1')
            top.set(qn('w:color'), '2E86AB')
            pBdr.append(top)
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single')
            bot.set(qn('w:sz'), '12')
            bot.set(qn('w:space'), '1')
            bot.set(qn('w:color'), 'F4A923')
            pBdr.append(bot)
            pPr.append(pBdr)
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(12) if level == 1 else Pt(10)
            run.font.color.rgb = RGB_BLAU
            run.font.name = 'Calibri'
            return p

        def add_taula_word(doc, dades, col_widths_cm, header=True):
            taula = doc.add_table(rows=len(dades), cols=len(dades[0]))
            taula.alignment = WD_TABLE_ALIGNMENT.LEFT
            for i, fila in enumerate(dades):
                for j, val in enumerate(fila):
                    cell = taula.rows[i].cells[j]
                    cell.width = Cm(col_widths_cm[j])
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(str(val))
                    run.font.name = 'Calibri'
                    if header and i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGB_BLANC
                        run.font.size = Pt(9)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_cell_bg(cell, '1A3A5C')
                        set_cell_border_bottom(cell, 'F4A923', 24)
                    else:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                        if i % 2 == 0:
                            set_cell_bg(cell, 'EEF2F7')
                        else:
                            set_cell_bg(cell, 'FFFFFF')
            return taula

        # ---- CAPÇALERA WORD ----
        # Linia blava superior
        p_top = doc_word.add_paragraph()
        pPr_top = p_top._p.get_or_add_pPr()
        pBdr_top = OxmlElement('w:pBdr')
        top_bdr = OxmlElement('w:top')
        top_bdr.set(qn('w:val'), 'single')
        top_bdr.set(qn('w:sz'), '36')
        top_bdr.set(qn('w:space'), '0')
        top_bdr.set(qn('w:color'), '1A3A5C')
        pBdr_top.append(top_bdr)
        pPr_top.append(pBdr_top)

        # Taula capçalera
        cap_w = doc_word.add_table(rows=1, cols=3)
        cap_w.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Col 1 — Nom empresa
        c0 = cap_w.rows[0].cells[0]
        c0.width = Cm(6)
        set_cell_bg(c0, 'EEF2F7')
        r0 = c0.paragraphs[0].add_run('ENERGYDATAGP')
        r0.font.bold = True; r0.font.size = Pt(14); r0.font.color.rgb = RGB_BLAU; r0.font.name = 'Calibri'
        c0.add_paragraph().add_run('Solucions d\'Analisi de Dades Energetiques').font.size = Pt(8)
        # Col 2 — Titol
        c1 = cap_w.rows[0].cells[1]
        c1.width = Cm(9)
        set_cell_bg(c1, 'EEF2F7')
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run('INFORME TECNIC DE PRODUCCIO\nINSTAL·LACIO FOTOVOLTAICA')
        r1.font.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = RGB_BLAU; r1.font.name = 'Calibri'
        # Col 3 — Ref
        c2 = cap_w.rows[0].cells[2]
        c2.width = Cm(4)
        set_cell_bg(c2, 'EEF2F7')
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(f'Ref.: {num_expedient if num_expedient else "-"}\nData: {datetime.now().strftime("%d/%m/%Y")}\nVersio: 1.0')
        r2.font.size = Pt(8); r2.font.color.rgb = RGB_BLAU; r2.font.name = 'Calibri'

        # Linia groga
        p_groc = doc_word.add_paragraph()
        pPr_g = p_groc._p.get_or_add_pPr()
        pBdr_g = OxmlElement('w:pBdr')
        bot_g = OxmlElement('w:bottom')
        bot_g.set(qn('w:val'), 'single')
        bot_g.set(qn('w:sz'), '24')
        bot_g.set(qn('w:space'), '0')
        bot_g.set(qn('w:color'), 'F4A923')
        pBdr_g.append(bot_g)
        pPr_g.append(pBdr_g)
        p_groc.paragraph_format.space_after = Pt(8)

        # ---- SECCIONS WORD ----
        add_heading_word(doc_word, '1. IDENTIFICACIO DEL PROJECTE')
        add_taula_word(doc_word, [
            ['Denominacio', nom_projecte, 'Promotor', nom_promotor if nom_promotor else '-'],
            ['Referencia', num_expedient if num_expedient else '-', 'Adreca', adreca if adreca else '-'],
            ['Coordenades', f'{lat}N, {lon}E', 'Any ref.', str(any_ref)],
            ['Tecnic', nom_tecnic, 'Col·legiacio', num_collegio if num_collegio else '-'],
        ], [4, 6.5, 3.5, 4.5], header=False)
        doc_word.add_paragraph()

        add_heading_word(doc_word, '2. ESTAT ACTUAL')
        p = doc_word.add_paragraph(text_estat)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

        add_heading_word(doc_word, '3. PROPOSTA DE MILLORA')
        p = doc_word.add_paragraph(text_proposta)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

        add_heading_word(doc_word, '4. PRESSUPOST DETALLAT')
        pres_rows = [['CONCEPTE', 'UNITATS', 'COST (EUR)']]
        for item in pressupost_final:
            cost_str = f"{item['cost']:,.2f}" if item['cost'] > 0 else '-'
            pres_rows.append([item['concepte'], str(item['unitats']), cost_str])
        pres_rows.append(['TOTAL (sense IVA)', '', f'{total_sense_iva:,.2f}'])
        pres_rows.append([f'TOTAL (amb IVA {iva}%)', '', f'{total_amb_iva:,.2f}'])
        t_pres = add_taula_word(doc_word, pres_rows, [10, 3, 4.5])
        # Total files en blau
        for ridx in [-2, -1]:
            for cell in t_pres.rows[ridx].cells:
                set_cell_bg(cell, 'D6E4F0')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGB_BLAU
        doc_word.add_paragraph()

        add_heading_word(doc_word, '5. PARAMETRES TECNICS')
        add_taula_word(doc_word, [
            ['Parametre', 'Valor', 'Unitat', 'Observacions'],
            ['Potencia pic', f'{potencia_kwp:.1f}', 'kWp', 'Condicions STC'],
            ['Nombre de moduls', f'{num_panels}', 'ut.', f'{potencia_panel_wp} Wp/modul'],
            ['Perdues', f'{perdues}', '%', 'Sistema complet'],
            ['Produccio anual', f'{produccio_anual:,.0f}', 'kWh/any', f'Any {any_ref}'],
            ['Mes maxim', mes_max, f'{produccio_mensual.max():,.0f} kWh', 'Pic estival'],
            ['Mes minim', mes_min, f'{produccio_mensual.min():,.0f} kWh', 'Minim hivernal'],
            ['HES/any', f'{produccio_anual/potencia_kwp:,.0f}', 'h', 'Hores a potencia nominal'],
            ['Orientacio/Inclinacio', orientacio, f'{inclinacio} graus', tipus_coberta],
        ], [5.5, 3, 3, 6])
        doc_word.add_paragraph()

        add_heading_word(doc_word, '6. PRODUCCIO MENSUAL')
        doc_word.add_picture('g1_mensual.png', width=Inches(6.2))
        doc_word.add_paragraph()

        add_heading_word(doc_word, '7. CORBES DE GENERACIO HORARIA')
        doc_word.add_picture('g2_corbes.png', width=Inches(6.2))
        doc_word.add_paragraph()
        doc_word.add_picture('g3_horaria.png', width=Inches(6.2))
        doc_word.add_paragraph()

        add_heading_word(doc_word, '8. ANALISI ECONOMICA')
        add_taula_word(doc_word, [
            ['Consum xarxa (kWh/any)', 'Generacio FV (kWh/any)', 'Autoconsum (kWh/any)', 'Excedents (kWh/any)'],
            [f'{consum_anual:,.0f}', f'{produccio_anual:,.0f}', f'{autoconsum_kwh:,.0f}', f'{excedents_kwh:,.0f}'],
        ], [4.5, 4.5, 4.5, 4.5])
        doc_word.add_paragraph()
        add_taula_word(doc_word, [
            ['Consum actual (kWh/any)', 'Estalvi energia (kWh/any)', 'Estalvi economic (EUR/any)', 'Inversio (EUR)', 'Payback (anys)', 'CO2 estalviat (kg/any)'],
            [f'{consum_anual:,.0f}', f'{autoconsum_kwh:,.0f}', f'{estalvi_anual:,.0f}', f'{cost_amb_iva:,.0f}', f'{anys_retorn:.1f}', f'{co2_estalviat:,.0f}'],
        ], [3, 3, 3.5, 3, 2.5, 3.5])
        doc_word.add_paragraph()
        doc_word.add_picture('g4_roi.png', width=Inches(6.2))
        doc_word.add_paragraph()

        add_heading_word(doc_word, '9. CONCLUSIONS')
        p = doc_word.add_paragraph(text_conclusions)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'
        doc_word.add_paragraph()

        # Signatura
        sig_t = doc_word.add_table(rows=3, cols=2)
        sig_t.alignment = WD_TABLE_ALIGNMENT.LEFT
        caps_sig = ['El/La tecnic/a responsable', 'Conforme el/la promotor/a']
        for j, cap in enumerate(caps_sig):
            c = sig_t.rows[0].cells[j]
            c.width = Cm(9)
            set_cell_bg(c, 'EEF2F7')
            set_cell_border_bottom(c, 'F4A923', 18)
            r = c.paragraphs[0].add_run(cap)
            r.font.bold = True; r.font.color.rgb = RGB_BLAU; r.font.size = Pt(9); r.font.name = 'Calibri'
        sig_t.rows[1].cells[0].width = Cm(9)
        sig_t.rows[1].cells[1].width = Cm(9)
        sig_t.rows[2].cells[0].paragraphs[0].add_run(f'Nom: {nom_tecnic}').font.size = Pt(9)
        sig_t.rows[2].cells[1].paragraphs[0].add_run('Nom:').font.size = Pt(9)
        doc_word.add_paragraph()

        # Peu Word
        p_peu = doc_word.add_paragraph()
        pPr_peu = p_peu._p.get_or_add_pPr()
        pBdr_peu = OxmlElement('w:pBdr')
        t_peu = OxmlElement('w:top')
        t_peu.set(qn('w:val'), 'single')
        t_peu.set(qn('w:sz'), '24')
        t_peu.set(qn('w:space'), '1')
        t_peu.set(qn('w:color'), '1A3A5C')
        pBdr_peu.append(t_peu)
        pPr_peu.append(pBdr_peu)
        r_peu = p_peu.add_run(f'EnergyDataGP — {datetime.now().strftime("%d/%m/%Y")} — PVGIS © Comissio Europea — www.energydatagp.com')
        r_peu.font.size = Pt(7); r_peu.font.color.rgb = RGB_BLAU_MIG; r_peu.font.name = 'Calibri'
        p_peu.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Avis legal Word
        p_av = doc_word.add_paragraph()
        r_av = p_av.add_run('AVIS LEGAL: Document generat amb EnergyDataGP amb finalitat orientativa. Les dades de produccio provenen de PVGIS © Comissio Europea. Aquest document no substitueix la memoria tecnica oficial signada per un tecnic collegiat competent.')
        r_av.font.size = Pt(7); r_av.font.color.rgb = RGBColor(0x88, 0x88, 0x88); r_av.font.name = 'Calibri'
        p_av.alignment = WD_ALIGN_PARAGRAPH.CENTER

        word_buffer = io.BytesIO()
        doc_word.save(word_buffer)
        word_buffer.seek(0)

    # =====================
    # EXCEL
    # =====================
    with st.spinner("Generant Excel..."):
        wb = openpyxl.Workbook()

        def estil_cap_excel(ws, row, cols, text_list, col_start=1):
            fill_blau = PatternFill("solid", fgColor="1A3A5C")
            font_blanc = Font(bold=True, color="FFFFFF", size=10, name='Calibri')
            border = Border(
                left=Side(style='thin', color='CCCCCC'),
                right=Side(style='thin', color='CCCCCC'),
                top=Side(style='thin', color='CCCCCC'),
                bottom=Side(style='medium', color='F4A923')
            )
            for i, text in enumerate(text_list):
                cell = ws.cell(row=row, column=col_start+i, value=text)
                cell.font = font_blanc
                cell.fill = fill_blau
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        def estil_fila_excel(ws, row, vals, col_start=1, bold_last=False):
            fill_gris = PatternFill("solid", fgColor="EEF2F7")
            fill_blanc = PatternFill("solid", fgColor="FFFFFF")
            border = Border(
                left=Side(style='thin', color='CCCCCC'),
                right=Side(style='thin', color='CCCCCC'),
                top=Side(style='thin', color='CCCCCC'),
                bottom=Side(style='thin', color='CCCCCC')
            )
            for i, val in enumerate(vals):
                cell = ws.cell(row=row, column=col_start+i, value=val)
                cell.fill = fill_gris if row % 2 == 0 else fill_blanc
                cell.border = border
                cell.font = Font(size=9, name='Calibri', bold=bold_last)
                cell.alignment = Alignment(horizontal='center', vertical='center')

        # ---- FULL 1: RESUM ----
        ws1 = wb.active
        ws1.title = "Resum"
        ws1.column_dimensions['A'].width = 35
        ws1.column_dimensions['B'].width = 22
        ws1.column_dimensions['C'].width = 15

        # Capçalera
        ws1.merge_cells('A1:C1')
        ws1['A1'] = 'ENERGYDATAGP — INFORME TECNIC FOTOVOLTAIC'
        ws1['A1'].font = Font(bold=True, size=14, color='1A3A5C', name='Calibri')
        ws1['A1'].fill = PatternFill("solid", fgColor="EEF2F7")
        ws1['A1'].alignment = Alignment(horizontal='center')
        ws1['A1'].border = Border(bottom=Side(style='medium', color='F4A923'))
        ws1.row_dimensions[1].height = 25

        ws1['A2'] = f'Projecte: {nom_projecte}'
        ws1['A3'] = f'Tecnic: {nom_tecnic} | Data: {datetime.now().strftime("%d/%m/%Y")}'
        for r in [2, 3]:
            ws1[f'A{r}'].font = Font(size=9, color='1A3A5C', name='Calibri')

        estil_cap_excel(ws1, 5, 3, ['PARAMETRE', 'VALOR', 'UNITAT'])
        dades_resum = [
            ('Potencia instal·lada', f'{potencia_kwp:.1f}', 'kWp'),
            ('Nombre de moduls', str(num_panels), 'ut.'),
            ('Produccio anual estimada', f'{produccio_anual:,.0f}', 'kWh/any'),
            ('Mes de maxima produccio', f'{mes_max} ({produccio_mensual.max():,.0f} kWh)', ''),
            ('Mes de minima produccio', f'{mes_min} ({produccio_mensual.min():,.0f} kWh)', ''),
            ('Hores equivalents de sol', f'{produccio_anual/potencia_kwp:,.0f}', 'HES/any'),
            ('Autoconsum', f'{autoconsum_kwh:,.0f}', 'kWh/any'),
            ('Excedents exportats', f'{excedents_kwh:,.0f}', 'kWh/any'),
            ('Estalvi economic anual', f'{estalvi_anual:,.0f}', 'EUR/any'),
            ('Cost instal·lacio amb IVA', f'{cost_amb_iva:,.0f}', 'EUR'),
            ('Payback', f'{anys_retorn:.1f}', 'anys'),
            ('CO2 estalviat', f'{co2_estalviat:,.0f}', 'kg/any'),
        ]
        for i, (p, v, u) in enumerate(dades_resum):
            estil_fila_excel(ws1, 6+i, [p, v, u])
            ws1.cell(row=6+i, column=1).alignment = Alignment(horizontal='left', vertical='center')

        # ---- FULL 2: PRODUCCIO MENSUAL ----
        ws2 = wb.create_sheet("Produccio Mensual")
        ws2.column_dimensions['A'].width = 15
        ws2.column_dimensions['B'].width = 22
        estil_cap_excel(ws2, 1, 2, ['MES', 'PRODUCCIO (kWh)'])
        mesos_noms = ['Gener', 'Febrer', 'Marc', 'Abril', 'Maig', 'Juny',
                     'Juliol', 'Agost', 'Setembre', 'Octubre', 'Novembre', 'Desembre']
        for i, (mes, prod) in enumerate(zip(mesos_noms, produccio_mensual.values)):
            estil_fila_excel(ws2, 2+i, [mes, round(prod, 0)])
        # Total
        ws2.cell(row=14, column=1, value='TOTAL ANUAL')
        ws2.cell(row=14, column=2, value=round(produccio_anual, 0))
        for col in [1, 2]:
            c = ws2.cell(row=14, column=col)
            c.font = Font(bold=True, color='1A3A5C', name='Calibri', size=10)
            c.fill = PatternFill("solid", fgColor="F4A923")
            c.border = Border(top=Side(style='medium', color='1A3A5C'))
            c.alignment = Alignment(horizontal='center')

        chart1 = BarChart()
        chart1.title = "Produccio Energetica Mensual (kWh)"
        chart1.style = 10
        chart1.y_axis.title = "kWh"
        chart1.x_axis.title = "Mes"
        chart1.width = 20; chart1.height = 12
        data_c1 = Reference(ws2, min_col=2, min_row=1, max_row=13)
        cats_c1 = Reference(ws2, min_col=1, min_row=2, max_row=13)
        chart1.add_data(data_c1, titles_from_data=True)
        chart1.set_categories(cats_c1)
        ws2.add_chart(chart1, "D2")

        # ---- FULL 3: CORBES HORARIES ----
        ws3 = wb.create_sheet("Corbes Horaries")
        ws3.column_dimensions['A'].width = 10
        mesos_corbes = [
            ('Gener (15 gen)', f'{any_ref}-01-15'),
            ('Abril (15 abr)', f'{any_ref}-04-15'),
            ('Juliol (15 jul)', f'{any_ref}-07-15'),
            ('Octubre (15 oct)', f'{any_ref}-10-15'),
            ('Estiu (21 jun)', f'{any_ref}-06-21'),
            ('Hivern (21 des)', f'{any_ref}-12-21'),
        ]
        cols_letters = ['B', 'C', 'D', 'E', 'F', 'G']
        for col_l in cols_letters:
            ws3.column_dimensions[col_l].width = 18
        estil_cap_excel(ws3, 1, 7, ['HORA'] + [m[0] for m in mesos_corbes])
        for h in range(24):
            row = h + 2
            ws3.cell(row=row, column=1, value=f'{h:02d}:00')
            fill = PatternFill("solid", fgColor="EEF2F7") if h % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
            ws3.cell(row=row, column=1).fill = fill
            ws3.cell(row=row, column=1).alignment = Alignment(horizontal='center')
            ws3.cell(row=row, column=1).font = Font(size=9, name='Calibri')
            for ci, (_, data) in enumerate(mesos_corbes):
                try:
                    dia_data = df.loc[data, 'P'] / 1000
                    vals_hora = dia_data[dia_data.index.hour == h]
                    val = round(float(vals_hora.mean()), 2) if len(vals_hora) > 0 else 0
                except:
                    val = 0
                c = ws3.cell(row=row, column=2+ci, value=val)
                c.fill = fill
                c.alignment = Alignment(horizontal='center')
                c.font = Font(size=9, name='Calibri')

        chart3 = LineChart()
        chart3.title = "Corbes de Generacio Horaria (kW)"
        chart3.style = 10
        chart3.y_axis.title = "Potencia (kW)"
        chart3.x_axis.title = "Hora del dia"
        chart3.width = 25; chart3.height = 15
        for ci in range(6):
            dr = Reference(ws3, min_col=2+ci, min_row=1, max_row=25)
            chart3.add_data(dr, titles_from_data=True)
        cats3 = Reference(ws3, min_col=1, min_row=2, max_row=25)
        chart3.set_categories(cats3)
        ws3.add_chart(chart3, "I2")

        # ---- FULL 4: PRESSUPOST ----
        ws4 = wb.create_sheet("Pressupost")
        ws4.column_dimensions['A'].width = 40
        ws4.column_dimensions['B'].width = 15
        ws4.column_dimensions['C'].width = 20
        estil_cap_excel(ws4, 1, 3, ['CONCEPTE', 'UNITATS', 'COST (EUR)'])
        for i, item in enumerate(pressupost_final):
            estil_fila_excel(ws4, 2+i, [item['concepte'], str(item['unitats']), item['cost'] if item['cost'] > 0 else '-'])
            ws4.cell(row=2+i, column=1).alignment = Alignment(horizontal='left', vertical='center')
            if item['cost'] > 0:
                ws4.cell(row=2+i, column=3).number_format = '#,##0.00 €'
        tr = len(pressupost_final)+2
        for ridx, (label, val) in enumerate([(f'TOTAL (sense IVA)', total_sense_iva), (f'TOTAL (amb IVA {iva}%)', total_amb_iva)]):
            r = tr + ridx
            ws4.cell(row=r, column=1, value=label)
            ws4.cell(row=r, column=3, value=val)
            for col in [1, 2, 3]:
                c = ws4.cell(row=r, column=col)
                c.font = Font(bold=True, color='1A3A5C', name='Calibri', size=10)
                c.fill = PatternFill("solid", fgColor="D6E4F0")
                c.border = Border(top=Side(style='medium', color='1A3A5C'))
            ws4.cell(row=r, column=3).number_format = '#,##0.00 €'
            ws4.cell(row=r, column=3).alignment = Alignment(horizontal='right')

        # ---- FULL 5: ROI ----
        ws5 = wb.create_sheet("ROI")
        ws5.column_dimensions['A'].width = 12
        ws5.column_dimensions['B'].width = 25
        ws5.column_dimensions['C'].width = 25
        estil_cap_excel(ws5, 1, 3, ['ANY', 'FLUX ANUAL (EUR)', 'FLUX ACUMULAT (EUR)'])
        for i in range(26):
            flux_an = estalvi_anual if i > 0 else -cost_amb_iva
            flux_ac = -cost_amb_iva + estalvi_anual * i
            estil_fila_excel(ws5, 2+i, [i, round(flux_an, 0), round(flux_ac, 0)])
            ws5.cell(row=2+i, column=2).number_format = '#,##0 €'
            ws5.cell(row=2+i, column=3).number_format = '#,##0 €'

        chart5 = LineChart()
        chart5.title = "Retorn de la Inversio (ROI)"
        chart5.style = 10
        chart5.y_axis.title = "EUR"
        chart5.x_axis.title = "Anys"
        chart5.width = 20; chart5.height = 12
        dr5 = Reference(ws5, min_col=3, min_row=1, max_row=27)
        chart5.add_data(dr5, titles_from_data=True)
        cats5 = Reference(ws5, min_col=1, min_row=2, max_row=27)
        chart5.set_categories(cats5)
        ws5.add_chart(chart5, "E2")

        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)

    # =====================
    # BOTONS DESCÀRREGA
    # =====================
    st.markdown("---")
    st.markdown("### 📥 Descarrega l'informe")
    nom_fitxer = num_expedient if num_expedient else nom_projecte.replace(' ', '_')
    col1, col2, col3 = st.columns(3)
    with col1:
        with open("informe_energydatagp.pdf", "rb") as f:
            st.download_button(
                label="📄 Descarregar PDF",
                data=f, file_name=f"informe_{nom_fitxer}.pdf",
                mime="application/pdf", use_container_width=True
            )
    with col2:
        st.download_button(
            label="📝 Descarregar Word (.docx)",
            data=word_buffer, file_name=f"informe_{nom_fitxer}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with col3:
        st.download_button(
            label="📊 Descarregar Excel (.xlsx)",
            data=excel_buffer, file_name=f"dades_{nom_fitxer}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )